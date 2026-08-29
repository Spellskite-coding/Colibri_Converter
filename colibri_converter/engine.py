"""
Moteur de conversion DOCX <-> PDF, hors-ligne.

Les fichiers d'entrée sont traités comme non fiables : voir SECURITY.md.
"""

from __future__ import annotations

import json
import logging
import os
import shutil
import signal
import subprocess
import tempfile
import time
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

from .render import worker as render_worker
from .render.isolation import WorkerFailed, apply_worker_limits, run_isolated

log = logging.getLogger("colibri_converter.engine")

# --------------------------------------------------------------------------
# Garde-fous de ressources
# --------------------------------------------------------------------------

MAX_INPUT_BYTES = 512 * 1024 * 1024        # refus au-delà : pas de DoS par taille
MAX_UNCOMPRESSED_BYTES = 2 * 1024**3       # zip-bomb : total décompressé
MAX_COMPRESSION_RATIO = 200                # zip-bomb : ratio par entrée
MAX_ZIP_ENTRIES = 10_000                   # zip-bomb : nombre d'entrées


class ConversionError(RuntimeError):
    """Erreur de conversion irrécupérable, message destiné à l'utilisateur."""


class BackendNotFound(ConversionError):
    """LibreOffice (ou une autre dépendance native) est introuvable."""


class UntrustedBackend(ConversionError):
    """Un binaire a été trouvé à un emplacement non fiable : on refuse de l'exécuter."""


@dataclass
class ConversionResult:
    source: Path
    output: Path
    backend: str
    duration_s: float = 0.0
    warnings: list[str] = field(default_factory=list)
    ocr_applied: bool = False


# --------------------------------------------------------------------------
# Résolution durcie d'un binaire externe (encore utilisé par ocrmypdf)
# --------------------------------------------------------------------------

# Emplacements inscriptibles par un utilisateur non privilégié : un binaire
# qui s'y trouve est soit un leurre, soit une install non maîtrisée.
_UNTRUSTED_MARKERS = frozenset({
    "downloads", "téléchargements", "temp", "tmp", "temporary internet files",
    "desktop", "bureau",
})


def _is_world_writable(path: Path) -> bool:
    """POSIX : un répertoire inscriptible par tous sur le chemin = vecteur de substitution."""
    if os.name == "nt":
        return False
    try:
        for parent in [path, *path.parents]:
            mode = parent.stat().st_mode if parent.exists() else 0
            if parent.exists() and (mode & 0o002) and not (mode & 0o1000):
                return True
    except OSError:
        return True
    return False


def _assert_trusted(binary: Path, *, origin: str) -> Path:
    """
    Refuse d'exécuter un binaire situé à un emplacement manipulable.
    C'est la contre-mesure au détournement de résolution (PATH / CWD hijacking).
    """
    try:
        resolved = binary.resolve(strict=True)
    except OSError as exc:
        raise BackendNotFound(f"Binaire inaccessible : {binary}") from exc

    if not resolved.is_file():
        raise UntrustedBackend(f"{resolved} n'est pas un fichier régulier.")

    # Jamais le répertoire courant : c'est le scénario de substitution classique.
    try:
        if resolved.parent == Path.cwd().resolve():
            raise UntrustedBackend(
                f"Binaire trouvé dans le répertoire courant ({resolved}) : refusé."
            )
    except OSError as exc:
        # Non bloquant : si le répertoire courant est illisible, on saute
        # simplement cette comparaison plutôt que de faire échouer toute
        # la résolution du binaire pour une vérification secondaire.
        log.debug("Comparaison au répertoire courant ignorée : %s", exc)

    # Comparaison par COMPOSANT de chemin, pas par sous-chaîne : un
    # utilisateur nommé "Templeton" contient "temp" et serait bloqué à tort.
    parts = {part.lower().strip() for part in resolved.parts}
    if parts & _UNTRUSTED_MARKERS:
        raise UntrustedBackend(
            f"Binaire situé dans un dossier temporaire ou de téléchargement ({resolved}) : "
            "refusé. Installe-le normalement plutôt que de le lancer depuis ce dossier."
        )

    if _is_world_writable(resolved.parent):
        raise UntrustedBackend(
            f"Le dossier {resolved.parent} est inscriptible par tous : exécution refusée."
        )

    log.debug("Backend retenu (%s) : %s", origin, resolved)
    return resolved


# --------------------------------------------------------------------------
# Exécution surveillée d'un sous-processus
# --------------------------------------------------------------------------


def _hardened_env() -> dict[str, str]:
    """Environnement d'un backend externe (ocrmypdf) : aucune sortie réseau possible."""
    env = dict(os.environ)
    for var in ("http_proxy", "https_proxy", "HTTP_PROXY", "HTTPS_PROXY",
                "ALL_PROXY", "all_proxy", "FTP_PROXY", "ftp_proxy"):
        env.pop(var, None)
    env["no_proxy"] = "*"
    env["NO_PROXY"] = "*"
    env["PYTHONDONTWRITEBYTECODE"] = "1"
    return env


def _kill_tree(proc: subprocess.Popen) -> None:
    """
    Détruit le processus ET sa descendance.
    Un simple kill() laisse un éventuel orphelin fuir en mémoire jusqu'au
    redémarrage de la session (ocrmypdf fork `tesseract` en sous-processus).
    """
    if proc.poll() is not None:
        return
    try:
        if os.name == "nt":
            taskkill = os.path.join(
                os.environ.get("SystemRoot", r"C:\\Windows"),
                "System32", "taskkill.exe",
            )
            subprocess.run(  # noqa: S603
                [taskkill, "/F", "/T", "/PID", str(proc.pid)],
                capture_output=True, timeout=15, check=False,
                creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
            )
        else:
            pgid = os.getpgid(proc.pid)
            os.killpg(pgid, signal.SIGTERM)
            # Attente active brève plutôt que proc.wait() : wait() se bloque
            # si le tube stdout est plein et que le processus écrit encore.
            deadline = time.monotonic() + 5
            while time.monotonic() < deadline and proc.poll() is None:
                time.sleep(0.05)
            if proc.poll() is None:
                os.killpg(pgid, signal.SIGKILL)
    except (OSError, subprocess.SubprocessError) as exc:
        log.warning("Nettoyage du processus %s incomplet : %s", proc.pid, exc)
    finally:
        deadline = time.monotonic() + 5
        while time.monotonic() < deadline and proc.poll() is None:
            time.sleep(0.05)
        if proc.poll() is None:
            log.error("Processus %s toujours vivant après kill.", proc.pid)


def _run_guarded(cmd: list[str], *, timeout: int, cwd: Path) -> tuple[int, str, str]:
    """
    Lance un backend natif sans shell, avec timeout et destruction d'arbre.
    `cwd` est un répertoire temporaire contrôlé : évite qu'une bibliothèque
    native charge une dépendance depuis un dossier utilisateur.
    """
    kwargs: dict = {
        "cwd": str(cwd),
        "env": _hardened_env(),
        "stdin": subprocess.DEVNULL,
        "stdout": subprocess.PIPE,
        "stderr": subprocess.PIPE,
        "text": True,
        "errors": "replace",
        "shell": False,
    }
    if os.name == "nt":
        kwargs["creationflags"] = (
            getattr(subprocess, "CREATE_NO_WINDOW", 0)
            | getattr(subprocess, "CREATE_NEW_PROCESS_GROUP", 0)
        )
    else:
        kwargs["start_new_session"] = True  # groupe de processus dédié pour killpg

    proc = subprocess.Popen(cmd, **kwargs)
    try:
        out, err = proc.communicate(timeout=timeout)
        return proc.returncode, out or "", err or ""
    except subprocess.TimeoutExpired:
        _kill_tree(proc)
        # Drainage obligatoire : si le backend a rempli le tube stdout, le
        # processus reste bloqué en écriture et n'est jamais récolté.
        try:
            proc.communicate(timeout=10)
        except (subprocess.TimeoutExpired, ValueError, OSError):
            log.warning("Tubes du processus %s non drainés.", proc.pid)
        raise ConversionError(
            f"Conversion interrompue après {timeout}s. "
            "Document trop lourd, ou fichier malformé."
        ) from None
    except BaseException:
        _kill_tree(proc)
        try:
            proc.communicate(timeout=10)
        except Exception as exc:
            # Best-effort : le processus est déjà tué, ce drainage ne sert
            # qu'à éviter un tube bloqué. Son échec ne doit rien changer
            # à l'exception d'origine qui remonte juste après.
            log.debug("Drainage des tubes après kill incomplet : %s", exc)
        raise
    finally:
        for stream in (proc.stdout, proc.stderr):
            if stream and not stream.closed:
                stream.close()


# --------------------------------------------------------------------------
# Validation d'entrée
# --------------------------------------------------------------------------


def _check_input(source: Path) -> Path:
    source = Path(source).expanduser()
    try:
        source = source.resolve(strict=True)
    except OSError as exc:
        raise ConversionError(f"Fichier introuvable : {source}") from exc
    if not source.is_file():
        raise ConversionError(f"Ce n'est pas un fichier : {source}")
    size = source.stat().st_size
    if size == 0:
        raise ConversionError(f"Fichier vide : {source.name}")
    if size > MAX_INPUT_BYTES:
        raise ConversionError(
            f"Fichier trop volumineux ({size / 1024**2:.0f} Mo, "
            f"limite {MAX_INPUT_BYTES / 1024**2:.0f} Mo)."
        )
    return source


def safe_output_path(source: Path, target_ext: str, outdir: Path | None = None) -> Path:
    """
    Chemin de sortie à côté du fichier source (ou dans outdir), sans jamais
    écraser un fichier existant. On n'utilise PAS le répertoire courant :
    en glisser-déposer sur une icône, il est imprévisible et parfois non
    inscriptible.
    """
    source = Path(source)
    directory = Path(outdir).expanduser().resolve() if outdir else source.parent
    directory.mkdir(parents=True, exist_ok=True)

    candidate = directory / f"{source.stem}{target_ext}"
    counter = 1
    while candidate.exists():
        candidate = directory / f"{source.stem} ({counter}){target_ext}"
        counter += 1
        if counter > 999:
            raise ConversionError("Trop de fichiers homonymes dans le dossier de sortie.")
    return candidate


# --------------------------------------------------------------------------
# Assainissement DOCX (macros, OLE, zip-bomb, zip-slip)
# --------------------------------------------------------------------------

# Types de relation dont une cible externe déclenche une connexion sortante
# ou un chargement de contenu actif au rendu. Les hyperliens sont conservés :
# ils ne sont suivis que sur clic, et les retirer dégraderait le document.
# Toujours neutralisés : une cible externe de ce type déclenche un
# chargement de contenu actif ou une exécution au rendu.
_RISKY_REL_TYPES = (
    "attachedTemplate", "oleObject", "package", "subDocument",
    "frame", "externalLink",
)

# Médias liés : neutralisés UNIQUEMENT si la cible est distante. Une image
# liée vers un fichier du disque est légitime et fréquente ; la supprimer
# faisait disparaître les illustrations du document converti.
_MEDIA_REL_TYPES = ("image", "audio", "video", "media")
_REMOTE_SCHEMES = ("http://", "https://", "ftp://", "ftps://", "//")

_RISKY_PARTS = ("word/vbaProject.bin", "word/vbaData.xml")
_RISKY_PREFIXES = ("word/embeddings/", "word/activeX/", "customUI/")


def _strip_external_rels(archive: zipfile.ZipFile, item, name: str) -> tuple[bytes, list[str]]:
    """
    Retire des fichiers .rels les relations pointant vers une cible EXTERNE
    et de type actif (gabarit distant, objet OLE, cadre...). Le moteur de
    rendu intégré ne suit déjà aucune de ces relations aujourd'hui, mais les
    retirer à la source reste la bonne défense en profondeur : ça évite
    qu'une évolution future du moteur (ou tout autre outil ouvrant cette
    copie assainie) ne les résolve par inadvertance.
    """
    # defusedxml bloque l'expansion d'entités au niveau du parseur, là où
    # xml.etree reste vulnérable. C'est la défense structurelle ; le
    # pré-filtre DTD/entité plus bas la double en profondeur.
    from defusedxml.ElementTree import fromstring as _defused_fromstring
    from xml.etree.ElementTree import ParseError, tostring

    # Un .rels légitime pèse quelques kilo-octets. Au-delà, on refuse de
    # parser : ElementTree reste sensible à l'expansion d'entités.
    if item.file_size > 1024 * 1024:
        raise ConversionError(f"Fichier de relations anormalement volumineux : {name}")

    limit = 1024 * 1024
    with archive.open(item) as handle:
        raw = handle.read(limit + 1)
    if len(raw) > limit:
        # L'en-tête mentait sur la taille : on refuse plutôt que de renvoyer
        # un contenu tronqué, qui corromprait le document réécrit.
        raise ConversionError(f"Fichier de relations tronqué ou piégé : {name}")

    # xml.etree est sensible à l'expansion d'entités (« billion laughs ») et à
    # l'expansion quadratique. Un .rels légitime n'a ni DTD ni entité : leur
    # présence suffit à rejeter le document, sans avoir à parser.
    head = raw[:4096].upper()
    if b"<!DOCTYPE" in head or b"<!ENTITY" in raw.upper():
        raise ConversionError(
            f"Déclaration XML interdite dans {name} (DTD ou entité). "
            "Conversion refusée."
        )

    try:
        root = _defused_fromstring(raw)
    except ParseError:
        return raw, [f"Relations illisibles, conservées telles quelles : {name}"]

    stripped: list[str] = []
    for child in list(root):
        if child.get("TargetMode") != "External":
            continue
        rel_type = (child.get("Type") or "").rsplit("/", 1)[-1]
        target = child.get("Target") or ""

        remote = target.lower().startswith(_REMOTE_SCHEMES)
        if rel_type in _RISKY_REL_TYPES or (rel_type in _MEDIA_REL_TYPES and remote):
            root.remove(child)
            stripped.append(
                f"Référence externe neutralisée ({rel_type}) : {target[:80]}"
            )

    if not stripped:
        return raw, []
    return tostring(root, encoding="UTF-8", xml_declaration=True), stripped


def _copy_bounded(reader, writer, remaining: int, name: str) -> int:
    """
    Copie par blocs en comptant les octets réellement produits.
    Ne fait jamais confiance à `file_size` de l'en-tête ZIP : ce champ est
    fourni par l'archive, donc par l'attaquant.
    """
    written = 0
    while True:
        chunk = reader.read(512 * 1024)
        if not chunk:
            return written
        written += len(chunk)
        if written > remaining:
            raise ConversionError(
                f"Archive suspecte : {name} dépasse le volume décompressé "
                "autorisé. Conversion refusée."
            )
        writer.write(chunk)


def sanitize_docx(source: Path, destination: Path) -> list[str]:
    """
    Réécrit le DOCX en retirant tout composant actif, et refuse les archives
    piégées. Retourne la liste des éléments neutralisés.
    """
    removed: list[str] = []
    total_uncompressed = 0

    try:
        with zipfile.ZipFile(source) as zin:
            entries = zin.infolist()
            if len(entries) > MAX_ZIP_ENTRIES:
                raise ConversionError(
                    f"Archive suspecte : {len(entries)} entrées "
                    f"(limite {MAX_ZIP_ENTRIES}). Conversion refusée."
                )

            with zipfile.ZipFile(destination, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in entries:
                    name = item.filename

                    # Zip-slip : un DOCX est un ZIP, donc un format d'archive hostile.
                    if name.startswith(("/", "\\")) or ".." in Path(name).parts or ":" in name:
                        removed.append(f"Entrée d'archive hors périmètre ignorée : {name}")
                        continue

                    if item.is_dir():
                        continue  # inutile en OOXML, et piégeux à réécrire

                    if name in _RISKY_PARTS or name.startswith(_RISKY_PREFIXES):
                        removed.append(f"Composant actif retiré : {name}")
                        continue

                    # Pré-filtre sur l'en-tête : élimine les cas grossiers
                    # sans décompresser. Ne remplace pas le comptage réel.
                    if item.compress_size > 0:
                        ratio = item.file_size / item.compress_size
                        if ratio > MAX_COMPRESSION_RATIO and item.file_size > 10 * 1024**2:
                            raise ConversionError(
                                f"Archive suspecte : taux de compression {ratio:.0f}:1 "
                                f"sur {name}. Conversion refusée."
                            )

                    # En-tête normalisé : une méthode de compression exotique
                    # ou chiffrée dans l'archive source ferait échouer l'écriture.
                    clean = zipfile.ZipInfo(name, date_time=item.date_time)
                    clean.compress_type = zipfile.ZIP_DEFLATED
                    clean.external_attr = item.external_attr

                    if name.endswith(".rels"):
                        payload, stripped = _strip_external_rels(zin, item, name)
                        removed += stripped
                        total_uncompressed += len(payload)
                        if total_uncompressed > MAX_UNCOMPRESSED_BYTES:
                            raise ConversionError(
                                "Archive suspecte : volume décompressé excessif."
                            )
                        zout.writestr(clean, payload)
                        continue

                    # Comptage réel des octets décompressés, jamais l'en-tête.
                    budget = MAX_UNCOMPRESSED_BYTES - total_uncompressed
                    with zin.open(item) as reader, zout.open(clean, "w") as writer:
                        total_uncompressed += _copy_bounded(reader, writer, budget, name)

    except zipfile.BadZipFile as exc:
        raise ConversionError(
            f"{Path(source).name} n'est pas un .docx valide (archive illisible)."
        ) from exc

    return removed


# --------------------------------------------------------------------------
# DOCX -> PDF
# --------------------------------------------------------------------------


def docx_to_pdf(
    source: Path,
    output: Path | None = None,
    *,
    pdfa: bool = False,
    timeout: int = 180,
    sanitize: bool = True,
) -> ConversionResult:
    """
    Convertit un .docx en PDF avec le moteur de rendu intégré (render/,
    basé sur ReportLab) : aucun programme externe, tout se passe dans un
    processus isolé de ce même exécutable.
    """
    source = _check_input(source)
    suffix = source.suffix.lower()
    if not suffix:
        raise ConversionError(
            f"{source.name} n'a pas d'extension : format indéterminable."
        )
    if suffix == ".pdf":
        raise ConversionError("Le fichier est déjà un PDF.")
    if suffix != ".docx":
        raise ConversionError(
            f"{source.name} : seul le format .docx est pris en charge par le "
            "moteur de rendu intégré. Convertis d'abord ce document en .docx "
            "(depuis Word ou LibreOffice), puis relance la conversion."
        )

    output = Path(output) if output else safe_output_path(source, ".pdf")
    warnings: list[str] = []
    started = time.monotonic()

    with tempfile.TemporaryDirectory(prefix="pdfconv_") as tmp:
        tmp_path = Path(tmp)
        # Nom neutralisé : le nom d'origine ne doit pas influencer le worker.
        work_src = tmp_path / "input.docx"

        if sanitize:
            warnings += sanitize_docx(source, work_src)
        else:
            shutil.copy2(source, work_src)

        produced = tmp_path / "output.pdf"
        try:
            warnings += render_worker.render_docx_isolated(
                work_src, produced, pdfa=pdfa, timeout=timeout,
            )
        except TimeoutError as exc:
            raise ConversionError(str(exc)) from None
        except RuntimeError as exc:
            raise ConversionError(
                f"Le document n'a pas pu être converti en PDF : {exc}\n"
                "Le fichier est peut-être corrompu ou utilise une "
                "fonctionnalité non prise en charge."
            ) from None

        output.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(str(produced), str(output))

    warnings.append(
        "PDF produit sans balisage structurel (accessibilité) : limitation "
        "actuelle du moteur de rendu intégré. Une conversion ultérieure "
        "PDF -> Word sera moins fidèle qu'avec un PDF balisé."
    )
    if pdfa:
        warnings.append(
            "Mode PDF/A best-effort (polices embarquées, transparence "
            "aplatie) : la conformité ISO 19005 n'est PAS vérifiée."
        )

    return ConversionResult(
        source=source, output=output,
        backend="colibri-render (ReportLab)",
        duration_s=time.monotonic() - started,
        warnings=warnings,
    )


# --------------------------------------------------------------------------
# PDF -> DOCX
# --------------------------------------------------------------------------


def _pdf2docx_worker(src: str, dst: str, pages: tuple[int, int] | None, sidecar: str) -> None:
    """
    Processus fils dédié. Le parsing PDF est la principale surface d'attaque
    du programme (PyMuPDF est du C++). Un crash ici ne touche pas l'application,
    et la mort du processus libère intégralement la mémoire — ce qui rend
    toute fuite du parseur sans conséquence sur la durée de vie de l'app.

    La récupération d'images manquantes (_recover_missing_images) tourne
    ICI, dans ce même processus isolé, et non après coup dans le parent :
    elle rouvre le PDF hostile avec PyMuPDF/Pillow, donc elle a exactement
    les mêmes exigences d'isolation que la conversion elle-même — la faire
    tourner sans filet dans le processus applicatif annulerait la garantie
    d'isolation pour ce chemin précis.
    """
    apply_worker_limits()
    from pdf2docx import Converter

    cv = Converter(src)
    try:
        if pages:
            cv.convert(dst, start=pages[0] - 1, end=pages[1])
        else:
            cv.convert(dst)
    finally:
        cv.close()

    recovered = _recover_missing_images(Path(src), Path(dst))
    Path(sidecar).write_text(json.dumps({"recovered_images": recovered}), encoding="utf-8")


def _pdf2docx_isolated(src: Path, dst: Path, pages, timeout: int) -> int:
    """
    Isolation factorisée dans render/isolation.py (voir render_worker pour
    le pendant DOCX -> PDF) : les deux directions parsent du contenu non
    fiable et partagent donc le même besoin de limite mémoire/timeout.
    Retourne le nombre d'images récupérées (voir _pdf2docx_worker).
    """
    sidecar = dst.with_suffix(dst.suffix + ".recovery.json")
    try:
        run_isolated(
            _pdf2docx_worker, (str(src), str(dst), pages, str(sidecar)),
            timeout=timeout, output_path=dst,
        )
    except WorkerFailed as exc:
        if exc.kind == "timeout":
            raise ConversionError(
                f"Reconstruction interrompue après {timeout}s. "
                "PDF trop complexe ou malformé."
            ) from None
        if exc.kind == "crash":
            raise ConversionError(
                "Échec de la reconstruction du document "
                f"(code {exc.detail}). Le PDF est probablement chiffré, "
                "corrompu, ou d'une structure non prise en charge."
            ) from None
        raise ConversionError("La reconstruction n'a produit aucun contenu exploitable.") from None

    recovered = 0
    if sidecar.is_file():
        try:
            recovered = json.loads(sidecar.read_text(encoding="utf-8")).get("recovered_images", 0)
        except (OSError, ValueError) as exc:
            log.debug("Compte d'images récupérées illisible : %s", exc)
        finally:
            sidecar.unlink(missing_ok=True)
    return recovered


def _text_density(pdf: Path) -> int | None:
    """Caractères extractibles par page. 0 => PDF image-only (scan)."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return None

    doc = None
    try:
        doc = fitz.open(pdf)
        if doc.is_encrypted and not doc.authenticate(""):
            raise ConversionError(
                "Ce PDF est protégé par mot de passe. Déchiffre-le en amont, "
                "avec les autorisations appropriées."
            )
        if doc.page_count == 0:
            return 0
        sampled = min(doc.page_count, 20)  # échantillon : évite de parser 5000 pages
        total = sum(len(doc[i].get_text("text").strip()) for i in range(sampled))
        return total // sampled
    except ConversionError:
        raise
    except Exception as exc:
        log.warning("Analyse de densité impossible : %s", exc)
        return None
    finally:
        if doc is not None:
            doc.close()  # libération explicite du handle natif


def _ocrmypdf() -> Path:
    found = shutil.which("ocrmypdf")
    if not found:
        raise BackendNotFound("ocrmypdf introuvable (requiert ocrmypdf + tesseract).")
    return _assert_trusted(Path(found), origin="PATH")


# --------------------------------------------------------------------------
# Récupération des images perdues par pdf2docx
# --------------------------------------------------------------------------

# En dessous de ce seuil (environ 20x20 pt), on considère l'image comme
# décorative (puce, filet, pixel de suivi) plutôt que comme un contenu perdu
# qui vaudrait la peine d'être récupéré.
_MIN_RECOVERABLE_IMAGE_AREA_PT2 = 400


def _pdf_significant_images(pdf_path: Path) -> list[bytes]:
    """
    Images 'significatives' du PDF source, tous xrefs uniques confondus.
    pdf2docx construit son propre modèle de mise en page à partir des blocs
    PyMuPDF et laisse parfois de côté des images (espaces colorimétriques,
    contenu vectoriel/clippé, groupes) : on repart directement de la source
    plutôt que de faire confiance à ce qu'il a choisi de garder.

    Deux passes, dans cet ordre :

    1. Par page (page.get_images) : donne accès à la géométrie de rendu, donc
       à l'aire réellement occupée — ce qui permet d'écarter les images
       décoratives (puces, filets) sous le seuil.

    2. Sur tout le document (doc.xref_length) : rattrape les images qui ne
       sont référencées par AUCUNE page. Un JPEG orphelin (présent comme
       objet mais absent des /Resources de chaque page — cas produit par
       certains générateurs de PDF) est invisible à la passe par page, donc
       à pdf2docx aussi ; sans cette seconde passe il serait perdu en
       silence. Faute de géométrie de page pour ces xrefs, on ne peut pas
       filtrer par aire : on retombe sur les dimensions en pixels pour
       écarter le vraiment minuscule.

    `seen_xrefs` est partagé entre les deux passes (et entre les pages) :
    une même image présente sur plusieurs pages, ou vue en passe 1 puis
    revue en passe 2, n'est extraite qu'une fois.
    """
    import fitz  # PyMuPDF

    images: list[bytes] = []
    seen_xrefs: set[int] = set()

    def _extract(doc, xref: int) -> None:
        if xref in seen_xrefs:
            return
        seen_xrefs.add(xref)
        try:
            info = doc.extract_image(xref)
        except Exception as exc:
            log.debug("Image xref %s non extractible : %s", xref, exc)
            return
        data = info.get("image") if info else None
        if data:
            images.append(data)

    # Seuil en pixels pour la passe 2 (pas de géométrie de page disponible).
    # ~20x20 px : équivalent grossier du seuil en points de la passe 1, pour
    # écarter un pixel de suivi ou une icône décorative orpheline.
    min_px_side = 20

    with fitz.open(pdf_path) as doc:
        # Passe 1 : par page, avec filtrage par aire de rendu.
        for page in doc:
            for img in page.get_images(full=True):
                xref = img[0]
                if xref in seen_xrefs:
                    continue
                try:
                    rects = page.get_image_rects(xref)
                    area = max((r.width * r.height for r in rects), default=0)
                except Exception:
                    area = 0
                if area and area < _MIN_RECOVERABLE_IMAGE_AREA_PT2:
                    seen_xrefs.add(xref)  # vue et écartée : ne pas la reprendre en passe 2
                    continue
                _extract(doc, xref)

        # Passe 2 : tout le document, pour les images non rattachées à une page.
        for xref in range(1, doc.xref_length()):
            if xref in seen_xrefs:
                continue
            try:
                if doc.xref_get_key(xref, "Subtype")[1] != "/Image":
                    continue
            except Exception:
                continue
            try:
                info = doc.extract_image(xref)
            except Exception as exc:
                log.debug("Image orpheline xref %s non extractible : %s", xref, exc)
                seen_xrefs.add(xref)
                continue
            data = info.get("image") if info else None
            if not data:
                seen_xrefs.add(xref)
                continue
            w, h = info.get("width", 0), info.get("height", 0)
            if w and h and (w < min_px_side or h < min_px_side):
                seen_xrefs.add(xref)
                continue
            seen_xrefs.add(xref)
            images.append(data)

    return images


def _image_pixel_size(data: bytes) -> tuple[int, int] | None:
    import io

    from PIL import Image

    try:
        with Image.open(io.BytesIO(data)) as im:
            return im.size
    except Exception:
        return None


def _docx_media_pixel_sizes(docx_path: Path) -> list[tuple[int, int]]:
    sizes: list[tuple[int, int]] = []
    with zipfile.ZipFile(docx_path) as z:
        for name in z.namelist():
            if not name.startswith("word/media/"):
                continue
            dims = _image_pixel_size(z.read(name))
            if dims:
                sizes.append(dims)
    return sizes


def _recover_missing_images(source_pdf: Path, docx_path: Path) -> int:
    """
    Compare les images significatives du PDF source à celles réellement
    présentes dans le .docx produit (par dimensions en pixels — pdf2docx
    réencode mais ne rééchantillonne pas les images qu'il garde) et rajoute
    en fin de document celles qui manquent, plutôt que de les perdre en
    silence. Position d'origine non garantie : mieux vaut une image présente
    au mauvais endroit qu'une image absente.
    """
    try:
        import io

        import docx
        from docx.shared import Inches
    except ImportError:
        return 0

    try:
        source_images = _pdf_significant_images(source_pdf)
    except Exception as exc:
        log.debug("Analyse des images source impossible : %s", exc)
        return 0
    if not source_images:
        return 0

    remaining_present = _docx_media_pixel_sizes(docx_path)
    missing: list[bytes] = []
    for data in source_images:
        dims = _image_pixel_size(data)
        if dims is None:
            continue
        match = next(
            (d for d in remaining_present if abs(d[0] - dims[0]) <= 2 and abs(d[1] - dims[1]) <= 2),
            None,
        )
        if match is not None:
            remaining_present.remove(match)
        else:
            missing.append(data)

    if not missing:
        return 0

    document = docx.Document(str(docx_path))
    document.add_paragraph(
        "--- Image(s) récupérée(s) automatiquement (position d'origine non garantie) ---"
    )
    recovered = 0
    max_width_in = 6.0
    for data in missing:
        try:
            dims = _image_pixel_size(data)
            kwargs = {}
            if dims and dims[0] / 96.0 > max_width_in:
                kwargs["width"] = Inches(max_width_in)
            document.add_picture(io.BytesIO(data), **kwargs)
            recovered += 1
        except Exception as exc:
            log.debug("Image récupérée mais non insérable dans le .docx : %s", exc)
    if recovered:
        document.save(str(docx_path))
    return recovered


def pdf_to_docx(
    source: Path,
    output: Path | None = None,
    *,
    ocr: str = "auto",
    ocr_lang: str = "fra+eng",
    pages: tuple[int, int] | None = None,
    timeout: int = 600,
) -> ConversionResult:
    """Reconstruit un .docx à partir d'un PDF (heuristique de mise en page)."""
    source = _check_input(source)
    output = Path(output) if output else safe_output_path(source, ".docx")
    started = time.monotonic()
    warnings: list[str] = []
    ocr_applied = False

    with tempfile.TemporaryDirectory(prefix="pdfconv_") as tmp:
        tmp_path = Path(tmp)
        working = tmp_path / "input.pdf"
        shutil.copy2(source, working)

        density = _text_density(working)
        if density is None:
            warnings.append("Analyse de la couche texte impossible.")
        elif ocr == "always" or (ocr == "auto" and density < 25):
            try:
                ocred = tmp_path / "ocr.pdf"
                code, _, err = _run_guarded(
                    [str(_ocrmypdf()), "--quiet", "--skip-text",
                     "--language", ocr_lang, "--optimize", "0",
                     str(working), str(ocred)],
                    timeout=timeout, cwd=tmp_path,
                )
                if code == 0 and ocred.is_file():
                    working, ocr_applied = ocred, True
                    warnings.append(
                        "Document scanné : OCR appliqué. "
                        "Relecture humaine indispensable."
                    )
                else:
                    warnings.append(f"OCR en échec : {(err or '').strip()[-200:]}")
            except ConversionError as exc:
                warnings.append(
                    f"Document scanné mais OCR indisponible ({exc}). "
                    "Le résultat contiendra surtout des images."
                )

        staging = tmp_path / "out.docx"
        recovered = _pdf2docx_isolated(working, staging, pages, timeout)
        if recovered:
            warnings.append(
                f"{recovered} image(s) absente(s) de la reconstruction ont "
                "été rajoutées en fin de document (position d'origine non "
                "garantie)."
            )

        output.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(str(staging), str(output))

    if not ocr_applied and density is not None and density < 25:
        warnings.append("Peu de texte détecté : la mise en page sera approximative.")

    return ConversionResult(
        source=source, output=output,
        backend="pdf2docx" + (" + OCR" if ocr_applied else ""),
        duration_s=time.monotonic() - started,
        warnings=warnings, ocr_applied=ocr_applied,
    )


# --------------------------------------------------------------------------

SUPPORTED_INPUT = {".docx", ".pdf"}


def convert(source: Path, output: Path | None = None, **kwargs) -> ConversionResult:
    """Aiguillage sur l'extension source."""
    ext = Path(source).suffix.lower()
    if ext == ".pdf":
        return pdf_to_docx(source, output, **kwargs)
    if ext in SUPPORTED_INPUT:
        return docx_to_pdf(source, output, **kwargs)
    raise ConversionError(
        f"Format non pris en charge : {ext or '(sans extension)'}. "
        f"Attendu : {', '.join(sorted(SUPPORTED_INPUT))}"
    )
