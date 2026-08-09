"""Mesure de la fidélité d'une conversion (texte + rendu visuel)."""

from __future__ import annotations

import logging

import difflib
import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

log = logging.getLogger("colibri_converter.validate")

# Seuils : en dessous, on alerte l'utilisateur au lieu de livrer en silence.
TEXT_OK = 0.98
TEXT_WARN = 0.90
VISUAL_OK = 0.97
VISUAL_WARN = 0.88


@dataclass
class FidelityReport:
    text_similarity: Optional[float] = None
    visual_similarity: Optional[float] = None
    page_count_source: Optional[int] = None
    page_count_output: Optional[int] = None
    issues: list[str] = field(default_factory=list)

    @property
    def verdict(self) -> str:
        scores = [s for s in (self.text_similarity, self.visual_similarity) if s is not None]
        if not scores:
            return "INDETERMINE"
        worst = min(scores)
        if self.issues and worst < TEXT_WARN:
            return "ECHEC"
        if worst >= TEXT_OK:
            return "CONFORME"
        if worst >= TEXT_WARN:
            return "DEGRADE"
        return "ECHEC"

    def summary(self) -> str:
        lines = [f"Verdict : {self.verdict}"]
        if self.text_similarity is not None:
            lines.append(f"  Fidélité textuelle : {self.text_similarity:6.2%}")
        if self.visual_similarity is not None:
            lines.append(f"  Fidélité visuelle  : {self.visual_similarity:6.2%}")
        if self.page_count_source is not None:
            lines.append(
                f"  Pages : {self.page_count_source} -> {self.page_count_output}"
            )
        for issue in self.issues:
            lines.append(f"  [!] {issue}")
        return "\n".join(lines)


# --------------------------------------------------------------------------
# Extraction de texte normalisé
# --------------------------------------------------------------------------


def _normalize(text: str) -> str:
    """
    Neutralise ce qui n'est pas une perte d'information réelle :
    ligatures, espaces insécables, césures de fin de ligne, blancs multiples.
    """
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\u00ad", "")                 # trait d'union conditionnel
    text = re.sub(r"-\s*\n\s*", "", text)             # césure typographique
    text = re.sub(r"[\s\u00a0\u2007\u202f]+", " ", text)
    return text.strip().lower()


def extract_text(path: Path) -> str:
    path = Path(path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        import fitz
        with fitz.open(path) as doc:
            return _normalize("\n".join(p.get_text("text") for p in doc))
    if ext == ".docx":
        import docx  # python-docx

        # python-docx décompresse sans plafond : on réutilise les garde-fous
        # du moteur plutôt que d'ouvrir l'archive brute.
        import tempfile

        from .engine import sanitize_docx

        with tempfile.TemporaryDirectory(prefix="colibri_audit_") as tmp:
            safe = Path(tmp) / "audit.docx"
            sanitize_docx(path, safe)
            return _extract_docx_text(safe)
    raise ValueError(f"Extraction non prise en charge pour {ext}")


def _extract_docx_text(path: Path) -> str:
    import docx

    try:
        d = docx.Document(str(path))
    except Exception as exc:
        # python-docx lève des KeyError/ValueError brutes sur un document
        # incomplet. On les traduit : l'audit ne doit jamais faire remonter
        # une trace interne à l'utilisateur.
        raise ValueError(f"Document illisible ou incomplet : {path.name}") from exc

    try:
        chunks = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                chunks.extend(c.text for c in row.cells)
        return _normalize("\n".join(chunks))
    except Exception as exc:
        raise ValueError(f"Extraction du texte impossible : {path.name}") from exc


# Au-delà de ce volume, difflib devient impraticable : sa complexité
# mémoire est quadratique, un document de 500 pages sature la RAM.
_DIFFLIB_LIMIT = 120_000


def _token_overlap(a: str, b: str) -> float:
    """
    Similarité par multiensemble de mots : linéaire en mémoire.
    Moins fine que difflib, mais c'est le seul choix viable sur les gros
    documents — et l'ordre des mots importe peu pour détecter une perte.
    """
    from collections import Counter

    ca, cb = Counter(a.split()), Counter(b.split())
    total = max(sum(ca.values()), sum(cb.values()))
    if not total:
        return 1.0
    return sum((ca & cb).values()) / total


def text_similarity(a: Path, b: Path) -> float:
    """Ratio de similarité 0..1 entre les contenus textuels normalisés."""
    ta, tb = extract_text(a), extract_text(b)
    if not ta and not tb:
        return 1.0
    if not ta or not tb:
        return 0.0
    if max(len(ta), len(tb)) > _DIFFLIB_LIMIT:
        return _token_overlap(ta, tb)
    return difflib.SequenceMatcher(None, ta, tb, autojunk=False).ratio()


# --------------------------------------------------------------------------
# Comparaison visuelle page à page
# --------------------------------------------------------------------------


def visual_similarity(pdf_a: Path, pdf_b: Path, dpi: int = 100) -> tuple[float, list[str]]:
    """
    Rend les deux PDF et compare page par page.
    Retourne (score moyen, anomalies). Utilise SSIM si scikit-image est
    disponible, sinon une distance pixel normalisée.
    """
    import fitz
    import numpy as np

    try:
        from skimage.metrics import structural_similarity as ssim
    except ImportError:
        ssim = None

    issues: list[str] = []
    scores: list[float] = []
    zoom = dpi / 72.0
    matrix = fitz.Matrix(zoom, zoom)
    MAX_PAGES = 200  # au-delà, on échantillonne : le rendu intégral est trop coûteux

    with fitz.open(pdf_a) as da, fitz.open(pdf_b) as db:
        if da.page_count != db.page_count:
            issues.append(
                f"Nombre de pages divergent : {da.page_count} vs {db.page_count} "
                "(repagination — vérifier polices et marges)."
            )
        comparable = min(da.page_count, db.page_count)
        if comparable > MAX_PAGES:
            issues.append(f"Comparaison visuelle limitée aux {MAX_PAGES} premières pages.")
            comparable = MAX_PAGES
        for i in range(comparable):
            ia = _render_gray(da[i], matrix, np)
            ib = _render_gray(db[i], matrix, np)
            h = min(ia.shape[0], ib.shape[0])
            w = min(ia.shape[1], ib.shape[1])
            ia, ib = ia[:h, :w], ib[:h, :w]

            if ssim is not None:
                score = float(ssim(ia, ib, data_range=255))
            else:
                score = 1.0 - float(np.abs(ia.astype(np.int16) - ib.astype(np.int16)).mean()) / 255.0

            scores.append(score)
            if score < VISUAL_WARN:
                issues.append(f"Page {i + 1} : divergence visuelle marquée ({score:.1%}).")

    return (sum(scores) / len(scores) if scores else 0.0), issues


def _render_gray(page, matrix, np):
    pix = page.get_pixmap(matrix=matrix, colorspace="gray", alpha=False)
    try:
        # copy() : frombuffer référence le tampon natif du pixmap. Sans copie,
        # l'image devient invalide dès la libération et peut lire de la
        # mémoire recyclée.
        return np.frombuffer(pix.samples, dtype=np.uint8).reshape(
            pix.height, pix.width
        ).copy()
    finally:
        pix = None


# --------------------------------------------------------------------------
# Point d'entrée
# --------------------------------------------------------------------------


def audit(source: Path, output: Path, *, reference_pdf: Optional[Path] = None) -> FidelityReport:
    """
    Compare la sortie à sa source. Si reference_pdf est fourni (aller-retour),
    on ajoute la comparaison visuelle.
    """
    report = FidelityReport()

    try:
        report.text_similarity = text_similarity(source, output)
    except Exception as exc:
        report.issues.append(f"Comparaison textuelle impossible : {exc}")

    if report.text_similarity is not None and report.text_similarity < TEXT_WARN:
        report.issues.append(
            "Perte de contenu significative : contrôle manuel indispensable "
            "(colonnes, notes de bas de page, tableaux imbriqués)."
        )

    if reference_pdf and Path(output).suffix.lower() == ".pdf":
        try:
            score, issues = visual_similarity(reference_pdf, output)
            report.visual_similarity = score
            report.issues.extend(issues)
        except Exception as exc:
            report.issues.append(f"Comparaison visuelle impossible : {exc}")

    try:
        import fitz
        for path, attr in ((source, "page_count_source"), (output, "page_count_output")):
            if Path(path).suffix.lower() == ".pdf":
                with fitz.open(path) as d:
                    setattr(report, attr, d.page_count)
    except Exception as exc:
        # Enrichissement facultatif du rapport : son échec n'invalide pas
        # l'audit lui-même, qui repose sur les similarités déjà calculées.
        log.debug("Comptage des pages ignoré : %s", exc)

    return report


def roundtrip_audit(docx_source: Path, workdir: Path) -> FidelityReport:
    """
    Test de non-régression le plus sévère : DOCX -> PDF -> DOCX -> PDF.
    Le premier et le dernier PDF doivent être visuellement quasi identiques.
    À câbler dans la CI sur un corpus de documents de référence.
    """
    from .engine import docx_to_pdf, pdf_to_docx

    workdir = Path(workdir)
    workdir.mkdir(parents=True, exist_ok=True)
    pdf1 = workdir / "step1.pdf"
    docx2 = workdir / "step2.docx"
    pdf2 = workdir / "step3.pdf"

    docx_to_pdf(docx_source, pdf1)
    pdf_to_docx(pdf1, docx2)
    docx_to_pdf(docx2, pdf2)

    return audit(pdf1, pdf2, reference_pdf=pdf1)
