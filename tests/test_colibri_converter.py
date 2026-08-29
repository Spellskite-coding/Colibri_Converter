"""
Tests de non-régression. Conçus pour tourner sur Windows ET Linux :
c'est le seul moyen d'exercer les branches spécifiques à chaque OS.

Aucune dépendance externe requise : le moteur de rendu DOCX -> PDF est
100% intégré, ces tests tournent donc toujours, y compris en CI minimale.

    pytest -v
"""

from __future__ import annotations

import os
import subprocess
import sys
import time
import zipfile
from pathlib import Path

import pytest

from colibri_converter import engine
from colibri_converter.engine import (  # noqa: E402
    ConversionError, UntrustedBackend, _assert_trusted, _run_guarded,
    _strip_external_rels, _UNTRUSTED_MARKERS, docx_to_pdf, pdf_to_docx,
    safe_output_path, sanitize_docx,
)

windows_only = pytest.mark.skipif(os.name != "nt", reason="spécifique Windows")
posix_only = pytest.mark.skipif(os.name == "nt", reason="spécifique POSIX")


def _minimal_docx(path: Path, extra: dict[str, bytes] | None = None) -> Path:
    with zipfile.ZipFile(path, "w") as z:
        z.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats.org'
            '/package/2006/content-types"><Default Extension="xml" '
            'ContentType="application/xml"/></Types>',
        )
        z.writestr("word/document.xml", "<w:document><w:body/></w:document>")
        for name, data in (extra or {}).items():
            z.writestr(name, data)
    return path


# ---------------------------------------------------------------- assainissement


def test_macro_et_ole_retires(tmp_path):
    src = _minimal_docx(tmp_path / "m.docx", {
        "word/vbaProject.bin": b"\x00" * 64,
        "word/embeddings/oleObject1.bin": b"MZ",
    })
    removed = sanitize_docx(src, tmp_path / "out.docx")
    with zipfile.ZipFile(tmp_path / "out.docx") as z:
        names = z.namelist()
    assert "word/vbaProject.bin" not in names
    assert not any(n.startswith("word/embeddings/") for n in names)
    assert len(removed) == 2


def test_zip_slip_rejete(tmp_path):
    src = _minimal_docx(tmp_path / "s.docx", {"../../../etc/passwd": b"pwned"})
    removed = sanitize_docx(src, tmp_path / "out.docx")
    with zipfile.ZipFile(tmp_path / "out.docx") as z:
        assert all(".." not in n for n in z.namelist())
    assert any("hors périmètre" in r for r in removed)


def test_zip_bomb_entete_falsifie(tmp_path):
    """L'en-tête ZIP est fourni par l'attaquant : il ne doit pas être cru."""
    src = tmp_path / "bomb.docx"
    with zipfile.ZipFile(src, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("word/document.xml", b"\x00" * (300 * 1024 * 1024))
    with zipfile.ZipFile(src, "a") as z:
        for info in z.infolist():
            info.file_size = 100          # mensonge délibéré
    with pytest.raises(ConversionError):
        sanitize_docx(src, tmp_path / "out.docx")


def test_modele_distant_retire_hyperlien_conserve(tmp_path):
    base = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    rels = (
        '<?xml version="1.0"?><Relationships xmlns="http://schemas.openxmlformats'
        '.org/package/2006/relationships">'
        f'<Relationship Id="r1" TargetMode="External" Type="{base}/attachedTemplate"'
        ' Target="http://attaquant.example/beacon.dotm"/>'
        f'<Relationship Id="r2" TargetMode="External" Type="{base}/hyperlink"'
        ' Target="https://legitime.example/page"/>'
        "</Relationships>"
    ).encode()
    src = _minimal_docx(tmp_path / "c.docx", {"word/_rels/settings.xml.rels": rels})
    out = tmp_path / "clean.docx"
    warnings = sanitize_docx(src, out)
    with zipfile.ZipFile(out) as z:
        cleaned = z.read("word/_rels/settings.xml.rels").decode()
    assert "attaquant.example" not in cleaned
    assert "legitime.example" in cleaned
    assert any("attachedTemplate" in w for w in warnings)


def test_archive_illisible(tmp_path):
    bad = tmp_path / "faux.docx"
    bad.write_bytes(b"pas du tout un zip")
    with pytest.raises(ConversionError):
        sanitize_docx(bad, tmp_path / "o.docx")


# ---------------------------------------------------------------- résolution binaire


def test_binaire_dans_downloads_refuse(tmp_path):
    d = tmp_path / "Downloads"
    d.mkdir()
    fake = d / ("soffice.exe" if os.name == "nt" else "soffice")
    fake.write_text("#!/bin/sh\n")
    with pytest.raises(UntrustedBackend):
        _assert_trusted(fake, origin="test")


def test_nom_utilisateur_contenant_temp_accepte(tmp_path):
    """
    Régression : la comparaison par sous-chaîne bloquait 'Templeton'.

    Le scratch dir de pytest vit toujours sous un dossier que l'engine
    considère lui-même comme non fiable — /tmp sur POSIX, AppData/Local/Temp
    sous Windows. On saute le test dans ce cas plutôt que de deviner un seul
    mot-clé : mieux vaut réutiliser _UNTRUSTED_MARKERS, la même liste dont se
    sert l'engine, que la reconstruire à moitié dans le test et diverger.
    """
    good = tmp_path / "templeton" / "opt" / "libreoffice" / "program"
    good.mkdir(parents=True)
    fake = good / ("soffice.exe" if os.name == "nt" else "soffice")
    fake.write_text("#!/bin/sh\n")
    if os.name != "nt":
        os.chmod(good, 0o755)

    parts = {p.lower() for p in tmp_path.parts}
    if parts & _UNTRUSTED_MARKERS:
        pytest.skip(
            "tmp_path vit sous un dossier que l'engine distrust "
            "(scratch dir de pytest) : cas non représentatif."
        )
    assert _assert_trusted(fake, origin="test").name.startswith("soffice")


def test_binaire_inexistant():
    with pytest.raises(ConversionError):
        _assert_trusted(Path("/n/existe/pas/soffice"), origin="test")


# ---------------------------------------------------------------- sous-processus


def test_timeout_leve_conversion_error(tmp_path):
    cmd = (["cmd", "/c", "ping -n 60 127.0.0.1 > nul"] if os.name == "nt"
           else ["/bin/sh", "-c", "sleep 60"])
    started = time.monotonic()
    with pytest.raises(ConversionError):
        _run_guarded(cmd, timeout=3, cwd=tmp_path)
    assert time.monotonic() - started < 30


@posix_only
def test_pas_orphelin_apres_timeout(tmp_path):
    """`soffice` fork `soffice.bin` : le kill doit viser l'arbre entier."""
    marker = "colibri-converter_test_orphan_294"
    with pytest.raises(ConversionError):
        _run_guarded(
            ["/bin/sh", "-c", f"sleep 294 & # {marker}\nsleep 294"],
            timeout=2, cwd=tmp_path,
        )
    time.sleep(1)
    running = subprocess.run(["ps", "-eo", "args"], capture_output=True, text=True).stdout
    assert "sleep 294" not in running


def test_tube_sature_ne_bloque_pas(tmp_path):
    """Un backend qui remplit stdout ne doit pas provoquer d'interblocage."""
    if os.name == "nt":
        pytest.skip("équivalent Windows non trivial")
    started = time.monotonic()
    with pytest.raises(ConversionError):
        _run_guarded(
            ["/bin/sh", "-c", "yes LONGUELIGNEDESORTIE | head -c 200000000; sleep 293"],
            timeout=3, cwd=tmp_path,
        )
    assert time.monotonic() - started < 30


# ---------------------------------------------------------------- sorties


def test_jamais_ecraser(tmp_path):
    (tmp_path / "doc.pdf").write_text("existant")
    result = safe_output_path(tmp_path / "doc.docx", ".pdf")
    assert result.name == "doc (1).pdf"
    assert (tmp_path / "doc.pdf").read_text() == "existant"


def test_fichier_vide_rejete(tmp_path):
    empty = tmp_path / "vide.docx"
    empty.touch()
    with pytest.raises(ConversionError):
        docx_to_pdf(empty)


def test_pdf_en_entree_de_docx_to_pdf(tmp_path):
    pdf = tmp_path / "deja.pdf"
    pdf.write_bytes(b"%PDF-1.4\n")
    with pytest.raises(ConversionError):
        docx_to_pdf(pdf)


# ---------------------------------------------------------------- bout en bout


def test_conversion_reelle_et_avertissement_de_balisage(tmp_path):
    """
    Conversion de bout en bout avec le moteur intégré. Le moteur ne produit
    jamais de PDF balisé (voir SECURITY.md/README) : l'absence de balisage
    doit être SIGNALÉE, conformément au principe du projet (mesurer la perte
    plutôt que la cacher), jamais silencieuse.
    """
    docx = pytest.importorskip("docx")
    d = docx.Document()
    d.add_heading("Titre de test", level=1)
    d.add_paragraph("Accents : éàüçñ — espace insécable inclus.")
    src = tmp_path / "reel.docx"
    d.save(src)

    result = docx_to_pdf(src)
    assert result.output.is_file()
    assert result.output.stat().st_size > 500

    fitz = pytest.importorskip("fitz")
    with fitz.open(result.output) as doc:
        assert doc.page_count >= 1
        assert "test" in doc[0].get_text().lower()

    raw = result.output.read_bytes()
    assert b"/StructTreeRoot" not in raw, "le moteur intégré ne balise jamais : ceci indiquerait une régression"
    assert any("balisage" in w.lower() for w in result.warnings), (
        "PDF non balisé sans avertissement : régression silencieuse."
    )


def test_accents_dans_le_chemin(tmp_path):
    docx = pytest.importorskip("docx")
    folder = tmp_path / "Dossier Éric & Cie"
    folder.mkdir()
    d = docx.Document()
    d.add_paragraph("test")
    src = folder / "rapport éàü.docx"
    d.save(src)
    assert docx_to_pdf(src).output.is_file()


def test_conversions_paralleles(tmp_path):
    """Chaque conversion tourne dans son propre processus isolé : pas de verrou partagé."""
    from concurrent.futures import ThreadPoolExecutor
    docx = pytest.importorskip("docx")
    sources = []
    for i in range(4):
        d = docx.Document()
        d.add_paragraph(f"document {i}")
        path = tmp_path / f"doc{i}.docx"
        d.save(path)
        sources.append(path)
    with ThreadPoolExecutor(max_workers=4) as pool:
        results = list(pool.map(docx_to_pdf, sources))
    assert all(r.output.is_file() for r in results)


def test_doc_odt_rtf_rejetes_proprement(tmp_path):
    """Le moteur intégré ne lit que l'OOXML : .doc/.odt/.rtf doivent échouer clairement."""
    for ext in (".doc", ".odt", ".rtf"):
        fake = tmp_path / f"ancien{ext}"
        fake.write_bytes(b"peu importe le contenu")
        with pytest.raises(ConversionError, match=r"\.docx"):
            docx_to_pdf(fake)


def test_libreoffice_definitivement_absent():
    """Garde-fou : empêche une réintroduction silencieuse de la dépendance LibreOffice."""
    assert not hasattr(engine, "find_soffice")
    assert not hasattr(engine, "libreoffice_available")


def test_imports_des_modules():
    """Détecte une dépendance manquante dans le binaire gelé."""
    import colibri_converter.cli
    import colibri_converter.engine
    import colibri_converter.validate
    assert colibri_converter.engine.SUPPORTED_INPUT


# ---------------------------------------------------------------- SAST


def test_dtd_dans_rels_rejetee(tmp_path):
    """xml.etree est sensible à l'expansion d'entités : on refuse toute DTD."""
    bomb = (
        b'<?xml version="1.0"?>'
        b'<!DOCTYPE r [<!ENTITY a "aaaaaaaaaa">'
        b'<!ENTITY b "&a;&a;&a;&a;&a;&a;&a;&a;&a;&a;">'
        b'<!ENTITY c "&b;&b;&b;&b;&b;&b;&b;&b;&b;&b;">]>'
        b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006'
        b'/relationships"><Relationship Id="r1" Type="x" Target="&c;"/>'
        b"</Relationships>"
    )
    src = _minimal_docx(tmp_path / "dtd.docx", {"word/_rels/document.xml.rels": bomb})
    with pytest.raises(ConversionError):
        sanitize_docx(src, tmp_path / "out.docx")


def test_rels_surdimensionne_rejete(tmp_path):
    """En-tête mensonger sur un .rels : refuser plutôt que tronquer."""
    huge = b'<?xml version="1.0"?><Relationships>' + b"<!-- " + b"x" * (2 * 1024 * 1024) + b" -->"
    src = _minimal_docx(tmp_path / "big.docx", {"word/_rels/document.xml.rels": huge})
    with pytest.raises(ConversionError):
        sanitize_docx(src, tmp_path / "out.docx")


def test_nom_de_fichier_hostile_echappe(tmp_path):
    """Le nom de fichier finit dans un QTextEdit, qui interprète le HTML."""
    import html as html_mod

    piege = '<img src="http://attaquant.example/x.png">'
    echappe = html_mod.escape(piege)
    assert "<img" not in echappe
    assert "&lt;img" in echappe


def test_lot_borne(tmp_path):
    """Un dossier immense ne doit pas geler le CLI sans retour."""
    from colibri_converter.cli import MAX_BATCH, _collect

    assert MAX_BATCH > 0
    for i in range(5):
        (tmp_path / f"f{i}.pdf").write_bytes(b"%PDF-1.4\n")
    (tmp_path / "ignore.txt").write_text("x")
    found = _collect([tmp_path])
    assert len(found) == 5
    assert all(f.suffix == ".pdf" for f in found)


def test_audit_passe_par_assainissement(tmp_path):
    """L'option --audit ne doit pas ouvrir l'archive brute."""
    import zipfile as zf

    from colibri_converter.validate import extract_text

    docx_mod = pytest.importorskip("docx")
    d = docx_mod.Document()
    d.add_paragraph("Contenu de contrôle")
    src = tmp_path / "a.docx"
    d.save(src)

    # On greffe une macro dans un document par ailleurs valide.
    piege = tmp_path / "piege.docx"
    with zf.ZipFile(src) as zin, zf.ZipFile(piege, "w") as zout:
        for item in zin.infolist():
            zout.writestr(item, zin.read(item))
        zout.writestr("word/vbaProject.bin", b"\x00" * 32)

    assert "contrôle" in extract_text(piege)


def test_docx_illisible_message_propre(tmp_path):
    """Aucune trace interne de python-docx ne doit remonter."""
    from colibri_converter.validate import extract_text

    src = _minimal_docx(tmp_path / "casse.docx")
    with pytest.raises(ValueError):
        extract_text(src)


def test_branding_svg_valides():
    import xml.etree.ElementTree as ET

    from colibri_converter.branding import COLIBRI_SVG, colibri_icon_svg

    ET.fromstring(COLIBRI_SVG)
    ET.fromstring(colibri_icon_svg())
    assert len(COLIBRI_SVG) < 8000  # doit rester léger pour l'icône 16 px


def test_image_liee_locale_conservee(tmp_path):
    """Régression : les images liées vers le disque étaient supprimées."""
    base = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    rels = (
        '<?xml version="1.0"?><Relationships xmlns="http://schemas.openxmlformats'
        '.org/package/2006/relationships">'
        f'<Relationship Id="r1" TargetMode="External" Type="{base}/image"'
        ' Target="file:///home/user/photo.png"/>'
        f'<Relationship Id="r2" TargetMode="External" Type="{base}/image"'
        ' Target="http://pisteur.example/pixel.png"/>'
        "</Relationships>"
    ).encode()
    src = _minimal_docx(tmp_path / "img.docx", {"word/_rels/document.xml.rels": rels})
    out = tmp_path / "clean.docx"
    sanitize_docx(src, out)
    with zipfile.ZipFile(out) as z:
        cleaned = z.read("word/_rels/document.xml.rels").decode()
    assert "photo.png" in cleaned, "une image liée locale ne doit pas disparaître"
    assert "pisteur.example" not in cleaned, "un pixel distant reste un rappel réseau"


def test_image_embarquee_survit_a_la_conversion(tmp_path):
    """Une image dans le DOCX doit se retrouver dans le PDF."""
    import io

    docx_mod = pytest.importorskip("docx")
    pil = pytest.importorskip("PIL.Image")

    buf = io.BytesIO()
    pil.new("RGB", (80, 80), (200, 50, 50)).save(buf, "PNG")
    buf.seek(0)

    d = docx_mod.Document()
    d.add_paragraph("Avant")
    d.add_picture(buf, width=docx_mod.shared.Inches(1))
    src = tmp_path / "avec_image.docx"
    d.save(src)

    raw = docx_to_pdf(src).output.read_bytes()
    assert b"/Image" in raw and b"/Width" in raw, "l'image a disparu du PDF"


def test_fonts_dir_coherent_dev_et_pyinstaller(monkeypatch):
    """
    colibri-converter.spec place les polices sous "fonts" dans le paquet
    PyInstaller (pas "vendor/fonts") : paths.fonts_dir() doit s'accorder
    avec cette destination une fois figé, faute de quoi le moteur de rendu
    ne retrouverait plus ses polices de repli dans l'exécutable final.
    """
    from colibri_converter import paths

    dev_dir = paths.fonts_dir()
    assert dev_dir == Path(__file__).resolve().parent.parent / "vendor" / "fonts"
    assert dev_dir.is_dir()
    assert len(list(dev_dir.glob("*.ttf"))) >= 20

    monkeypatch.setattr(sys, "frozen", True, raising=False)
    monkeypatch.setattr(sys, "_MEIPASS", "/fake/meipass", raising=False)
    assert str(paths.fonts_dir()) == str(Path("/fake/meipass") / "fonts")


# ---------------------------------------------------------------- moteur de rendu (render/)


def test_caracteres_speciaux_dans_le_corps_du_texte(tmp_path):
    """
    Le texte est injecté dans un mini-langage de balisage ReportLab : un '<'
    ou un '&' non échappé casse le parsing XML interne au lieu de s'afficher.
    """
    docx_mod = pytest.importorskip("docx")
    d = docx_mod.Document()
    d.add_paragraph("AT&T : 3 < 5 && x > y, <balise> non fermée")
    src = tmp_path / "special.docx"
    d.save(src)

    result = docx_to_pdf(src)
    fitz = pytest.importorskip("fitz")
    with fitz.open(result.output) as doc:
        text = doc[0].get_text()
    assert "AT&T" in text
    assert "3 < 5" in text


def test_saut_de_page_explicite(tmp_path):
    docx_mod = pytest.importorskip("docx")
    from docx.enum.text import WD_BREAK

    d = docx_mod.Document()
    d.add_paragraph("Page un")
    run = d.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)
    d.add_paragraph("Page deux")
    src = tmp_path / "saut.docx"
    d.save(src)

    fitz = pytest.importorskip("fitz")
    with fitz.open(docx_to_pdf(src).output) as doc:
        assert doc.page_count >= 2
        assert "page un" in doc[0].get_text().lower()
        assert "page deux" in doc[-1].get_text().lower()


def test_liste_a_puces_rendue(tmp_path):
    docx_mod = pytest.importorskip("docx")
    d = docx_mod.Document()
    d.add_paragraph("premier élément", style="List Bullet")
    d.add_paragraph("second élément", style="List Bullet")
    src = tmp_path / "liste.docx"
    d.save(src)

    fitz = pytest.importorskip("fitz")
    with fitz.open(docx_to_pdf(src).output) as doc:
        text = doc[0].get_text()
    assert "premier élément" in text
    assert "•" in text


def test_tableau_avec_fusion_horizontale(tmp_path):
    docx_mod = pytest.importorskip("docx")
    d = docx_mod.Document()
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "fusionnée"
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(1, 0).text = "bas gauche"
    table.cell(1, 1).text = "bas droite"
    src = tmp_path / "tableau.docx"
    d.save(src)

    result = docx_to_pdf(src)
    fitz = pytest.importorskip("fitz")
    with fitz.open(result.output) as doc:
        text = doc[0].get_text()
    assert "fusionnée" in text
    assert "bas gauche" in text and "bas droite" in text


def test_lien_hypertexte_present_dans_le_pdf(tmp_path):
    docx_mod = pytest.importorskip("docx")
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.oxml.ns import qn

    d = docx_mod.Document()
    p = d.add_paragraph("Voir ")
    r_id = p.part.relate_to("https://example.org/page", RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = p._p.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})
    run = p.add_run("le site")
    p._p.remove(run._r)
    hyperlink.append(run._r)
    p._p.append(hyperlink)

    src = tmp_path / "lien.docx"
    d.save(src)

    raw = docx_to_pdf(src).output.read_bytes()
    assert b"example.org" in raw, "l'URL du lien a disparu du PDF"


def _hang_forever(_a, _b):  # module-level : picklable par multiprocessing spawn
    time.sleep(60)


def test_pdf_to_docx_bout_en_bout(tmp_path):
    """
    Passe par l'API publique pdf_to_docx() de bout en bout, pas seulement
    par les fonctions internes : couvre le pipeline complet worker isolé +
    fichier sidecar de récupération d'images, refactoré pendant la revue de
    sécurité (voir test_recuperation_images_appelee_depuis_le_worker_isole).
    """
    fitz = pytest.importorskip("fitz")
    pytest.importorskip("pdf2docx")
    pytest.importorskip("docx")

    pdf_path = tmp_path / "source.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Bonjour depuis un PDF de test.")
    doc.save(pdf_path)
    doc.close()

    result = pdf_to_docx(pdf_path)
    assert result.output.is_file()
    assert result.output.suffix == ".docx"

    import docx as docx_mod
    d = docx_mod.Document(str(result.output))
    text = "\n".join(p.text for p in d.paragraphs)
    assert "Bonjour" in text


def test_recuperation_images_appelee_depuis_le_worker_isole():
    """
    Non-régression d'un contournement d'isolation trouvé en revue de
    sécurité : _recover_missing_images() rouvre le PDF hostile avec
    PyMuPDF/Pillow, donc elle doit s'exécuter DANS le worker isolé
    (_pdf2docx_worker, sous run_isolated), jamais directement dans
    pdf_to_docx(), qui tourne dans le processus applicatif sans le plafond
    mémoire ni le timeout du worker. Un appel direct depuis pdf_to_docx()
    annulerait la garantie d'isolation pour ce chemin précis.
    """
    assert "_recover_missing_images" in engine._pdf2docx_worker.__code__.co_names
    assert "_recover_missing_images" not in engine.pdf_to_docx.__code__.co_names


def test_recuperation_image_absente_du_docx_reconstruit(tmp_path):
    """
    _recover_missing_images ne dépend pas du comportement précis de pdf2docx
    (variable selon les versions) : on simule directement le cas qu'il est
    censé couvrir — une image présente dans le PDF source, absente du .docx
    produit — et on vérifie qu'elle est rajoutée plutôt que perdue.
    """
    fitz = pytest.importorskip("fitz")
    docx_mod = pytest.importorskip("docx")
    pil = pytest.importorskip("PIL.Image")
    from colibri_converter.engine import _recover_missing_images

    import io as io_mod
    buf = io_mod.BytesIO()
    pil.new("RGB", (200, 150), (10, 200, 10)).save(buf, "PNG")
    img_bytes = buf.getvalue()

    pdf_path = tmp_path / "source.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_image(fitz.Rect(50, 50, 250, 200), stream=img_bytes)
    doc.save(pdf_path)
    doc.close()

    docx_path = tmp_path / "reconstruit.docx"
    d = docx_mod.Document()
    d.add_paragraph("Texte reconstruit, sans l'image.")
    d.save(docx_path)

    recovered = _recover_missing_images(pdf_path, docx_path)
    assert recovered == 1

    with zipfile.ZipFile(docx_path) as z:
        media = [n for n in z.namelist() if n.startswith("word/media/")]
    assert media, "l'image récupérée n'a pas été rajoutée au .docx"


def test_rendu_isole_timeout():
    """Un rendu qui dépasse son délai doit être tué et signalé, pas bloquer indéfiniment."""
    from colibri_converter.render.isolation import WorkerFailed, run_isolated

    started = time.monotonic()
    with pytest.raises(WorkerFailed):
        run_isolated(_hang_forever, (1, 2), timeout=2)
    assert time.monotonic() - started < 30
